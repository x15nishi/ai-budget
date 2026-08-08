# AI Personal Budget Planner
# Mini Project - Python + Streamlit + Gemini API (via LangChain)
# made this to track monthly budget and get AI tips

import os
import io
import re
import json
import base64
import urllib.parse
from datetime import datetime

import pandas as pd
import streamlit as st
import plotly.express as px
import plotly.graph_objects as go
from dotenv import load_dotenv
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser
from langchain_core.messages import HumanMessage

load_dotenv()

DATA_FILE = "data.csv"
cols = ["Category", "Remark", "Amount", "Date"]  # columns for the expense table
CATEGORY_OPTIONS = ["Food", "Rent", "Transport", "Shopping", "Entertainment", "Utilities", "Other"]
INCOME_SOURCE_OPTIONS = ["Salary", "Interest", "Refund", "Dividend", "Other Income"]

# one consistent colour per category, used across the pie/bar charts and the budget bars
CATEGORY_COLORS = {
    "Food": "#F97316",
    "Rent": "#6366F1",
    "Transport": "#0EA5E9",
    "Shopping": "#EC4899",
    "Entertainment": "#8B5CF6",
    "Utilities": "#14B8A6",
    "Other": "#64748B",
}

st.set_page_config(page_title="AI Personal Budget Planner", layout="wide", page_icon="💸")

# =========================================================
# GLOBAL STYLE
# =========================================================
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700;800&family=Poppins:wght@600;700;800&display=swap');

html, body, [class*="css"]  { font-family: 'Inter', sans-serif; }

.main .block-container { padding-top: 1.6rem; padding-bottom: 3rem; max-width: 1200px; }

/* ---- hero header ---- */
.hero {
    background: linear-gradient(135deg, #4F46E5 0%, #7C3AED 100%);
    padding: 1.8rem 2.2rem;
    border-radius: 20px;
    color: white;
    margin-bottom: 1.4rem;
    box-shadow: 0 10px 30px rgba(79,70,229,0.25);
}
.hero h1 { font-family: 'Poppins', sans-serif; font-size: 1.7rem; margin: 0 0 4px 0; font-weight: 700; }
.hero p { margin: 0; opacity: 0.9; font-size: 0.92rem; }

/* ---- metric cards ---- */
.metric-row { display:flex; gap:14px; margin-bottom: 1.3rem; flex-wrap: wrap; }
.metric-card {
    flex:1; min-width: 190px;
    background: #FFFFFF;
    border-radius: 16px;
    padding: 1rem 1.25rem;
    box-shadow: 0 2px 10px rgba(15,23,42,0.06);
    border-left: 5px solid var(--accent, #4F46E5);
}
.metric-card .m-label { font-size: 0.75rem; color:#64748B; font-weight:700; text-transform:uppercase; letter-spacing:.05em; }
.metric-card .m-value { font-size: 1.45rem; font-weight:800; color:#0F172A; margin-top:2px; font-family:'Poppins',sans-serif; }
.metric-card .m-sub { font-size:0.76rem; color:#94A3B8; margin-top:3px; }

/* ---- generic section card ---- */
.section-card {
    background:#fff; border-radius:18px; padding:1.3rem 1.4rem;
    box-shadow:0 2px 10px rgba(15,23,42,0.05); margin-bottom:1.1rem;
}
.section-title {
    font-family:'Poppins', sans-serif; font-weight:700; font-size:1.02rem;
    color:#1E293B; margin-bottom: .7rem; display:flex; align-items:center; gap:8px;
}
.section-caption { color:#94A3B8; font-size:0.82rem; margin-top:-8px; margin-bottom:.8rem; }

/* ---- buttons ---- */
.stButton>button {
    border-radius: 10px !important;
    font-weight: 600 !important;
    transition: all .15s ease;
}
.stButton>button:hover { transform: translateY(-1px); box-shadow: 0 6px 14px rgba(79,70,229,0.20); }

/* ---- tabs ---- */
.stTabs [data-baseweb="tab-list"] { gap: 4px; }
.stTabs [data-baseweb="tab"] { border-radius: 10px 10px 0 0; padding: 10px 16px; font-weight:600; }

/* ---- custom budget progress bars ---- */
.budget-bar-wrap { margin-bottom: 14px; }
.budget-bar-label { display:flex; justify-content:space-between; font-size:0.83rem; font-weight:600; color:#334155; margin-bottom:4px; }
.budget-bar-track { background:#EEF2F7; border-radius: 8px; height:10px; overflow:hidden; }
.budget-bar-fill { height:100%; border-radius:8px; }

/* ---- badges ---- */
.badge { display:inline-block; padding:3px 10px; border-radius:999px; font-size:0.72rem; font-weight:700; }

/* ---- empty state ---- */
.empty-state { text-align:center; padding: 2.4rem 1rem; color:#94A3B8; }
.empty-state .icon { font-size:2.3rem; margin-bottom: 6px; }
.empty-state .title { color:#475569; font-weight:600; margin-bottom:2px; }

/* ---- sidebar quick stats ---- */
.side-stat { display:flex; justify-content:space-between; font-size:0.85rem; padding:6px 0; border-bottom:1px solid #EEF2F7; }
.side-stat b { color:#0F172A; }
</style>
""", unsafe_allow_html=True)


# =========================================================
# SMALL RENDER HELPERS
# =========================================================
def money(v):
    return f"₹{v:,.2f}"


def metric_card_html(label, value, sub, accent):
    # NOTE: no leading whitespace on these lines - Streamlit's markdown parser
    # treats 4+ space indented lines as a code block, which breaks the HTML render
    return (
        f'<div class="metric-card" style="--accent:{accent}">'
        f'<div class="m-label">{label}</div>'
        f'<div class="m-value">{value}</div>'
        f'<div class="m-sub">{sub}</div>'
        f'</div>'
    )


def render_summary_bar(income, total_exp, balance, savings):
    rate = (savings / income * 100) if income > 0 else 0
    cards = [
        metric_card_html("Income", money(income), "this month", "#4F46E5"),
        metric_card_html("Expense", money(total_exp), f"{len(st.session_state.expenses)} entries", "#F97316"),
        metric_card_html("Balance", money(balance), "income − expense", "#0EA5E9" if balance >= 0 else "#EF4444"),
        metric_card_html("Savings", money(savings), f"{rate:.1f}% of income" if income > 0 else "add income", "#22C55E" if rate >= 20 else "#F59E0B"),
    ]
    st.markdown(f'<div class="metric-row">{"".join(cards)}</div>', unsafe_allow_html=True)


def empty_state(icon, title, sub):
    st.markdown(
        f'<div class="empty-state">'
        f'<div class="icon">{icon}</div>'
        f'<div class="title">{title}</div>'
        f'<div>{sub}</div>'
        f'</div>',
        unsafe_allow_html=True
    )


def budget_bar(cat, spent, lim):
    pct = min(spent / lim, 1.0) * 100 if lim > 0 else 0
    color = "#22C55E" if pct < 80 else ("#F59E0B" if pct <= 100 else "#EF4444")
    over = spent > lim
    st.markdown(
        f'<div class="budget-bar-wrap">'
        f'<div class="budget-bar-label">'
        f'<span>{cat}</span>'
        f'<span>{money(spent)} / {money(lim)} {"⚠️" if over else ""}</span>'
        f'</div>'
        f'<div class="budget-bar-track">'
        f'<div class="budget-bar-fill" style="width:{pct}%; background:{color};"></div>'
        f'</div>'
        f'</div>',
        unsafe_allow_html=True
    )
    if over:
        st.markdown(f'<span class="badge" style="background:#FEE2E2;color:#B91C1C;">{cat} is over budget</span>', unsafe_allow_html=True)


# small helper so we don't repeat the same ChatGoogleGenerativeAI setup everywhere
def get_gemini_model():
    return ChatGoogleGenerativeAI(
        model="gemini-3.5-flash-lite",
        google_api_key=GEMINI_API_KEY,
        temperature=0.7
    )


# loads expense data from csv, if file not there just make empty table
def load_expenses():
    try:
        df = pd.read_csv(DATA_FILE)
        # older data.csv files (before the Date column was added) only had 3 cols -
        # just add the missing column instead of nuking their saved data
        if "Date" not in df.columns:
            df["Date"] = pd.NaT
        missing = [c for c in cols if c not in df.columns]
        if missing or set(df.columns) - set(cols):
            df = df.reindex(columns=cols)
    except Exception:
        df = pd.DataFrame(columns=cols)
    df["Amount"] = pd.to_numeric(df["Amount"], errors="coerce").fillna(0.0)
    df["Date"] = pd.to_datetime(df["Date"], errors="coerce")
    return df


def save_expenses(df):
    df.to_csv(DATA_FILE, index=False)


# ---------------- OCR BILL SCAN (Gemini Vision) ----------------
def extract_expenses_from_bill(image_bytes, mime_type):
    model = get_gemini_model()
    b64_img = base64.b64encode(image_bytes).decode("utf-8")

    instructions = (
        "You are reading a photo of a shopping receipt / bill / invoice. "
        "Extract each distinct expense line you can clearly identify. "
        "Respond with ONLY valid JSON (no markdown fences, no explanation text), "
        "as a JSON array like this exact shape:\n"
        '[{"category": "Food", "remark": "short item or store name", "amount": 123.45}]\n'
        "Rules:\n"
        f"- category MUST be exactly one of: {', '.join(CATEGORY_OPTIONS)}\n"
        "- amount must be a plain number (no currency symbols, no commas)\n"
        "- if you can only make out one grand total (not itemised), return a single "
        "object using the store/vendor name as remark and the total as amount\n"
        "- if the image is unreadable or not a bill, return an empty array []"
    )

    message = HumanMessage(
        content=[
            {"type": "text", "text": instructions},
            {"type": "image_url", "image_url": f"data:{mime_type};base64,{b64_img}"},
        ]
    )

    response = model.invoke([message])

    # response.content can be a plain string OR a list of content blocks
    # (e.g. [{"type": "text", "text": "..."}]) depending on the model/response,
    # so normalize it to a string first
    raw_content = response.content
    if isinstance(raw_content, list):
        text_parts = []
        for block in raw_content:
            if isinstance(block, dict):
                text_parts.append(block.get("text", ""))
            else:
                text_parts.append(str(block))
        raw_text = "".join(text_parts).strip()
    else:
        raw_text = str(raw_content).strip()

    if not raw_text:
        return []

    if raw_text.startswith("```"):
        raw_text = raw_text.strip("`")
        raw_text = raw_text.replace("json", "", 1).strip()

    items = json.loads(raw_text)

    cleaned = []
    for item in items:
        cat = item.get("category", "Other")
        if cat not in CATEGORY_OPTIONS:
            cat = "Other"
        try:
            amt = float(item.get("amount", 0))
        except (TypeError, ValueError):
            amt = 0.0
        remark = str(item.get("remark", "")).strip()
        if amt > 0:
            cleaned.append({"Category": cat, "Remark": remark, "Amount": amt})

    return cleaned


# =========================================================
# BANK STATEMENT IMPORT - parsing, EDA & auto-categorization
# (supports csv, xlsx/xls, pdf, docx, xml, txt - any format a bank export
# might come in, on a best-effort basis)
# =========================================================

# column name synonyms used to auto-detect which column is which, since every
# bank names things differently (e.g. HDFC uses "Narration", SBI uses "Description")
DATE_SYNONYMS = ["date", "txn date", "transaction date", "value date", "posting date", "tran date"]
DESC_SYNONYMS = ["description", "narration", "particulars", "details", "remarks", "transaction details", "reference", "remark"]
DEBIT_SYNONYMS = ["debit", "withdrawal", "withdrawal amt", "withdrawal amount", "dr", "paid out", "amount debited", "debit amount"]
CREDIT_SYNONYMS = ["credit", "deposit", "deposit amt", "deposit amount", "cr", "paid in", "amount credited", "credit amount"]
BALANCE_SYNONYMS = ["balance", "closing balance", "available balance", "running balance"]
AMOUNT_SYNONYMS = ["amount", "txn amount", "transaction amount", "amt"]
TYPE_SYNONYMS = ["type", "dr/cr", "cr/dr", "transaction type", "indicator"]

SELF_TRANSFER_KEYWORDS = ["self", "own a/c", "own account", "internal transfer", "to self", "imps-self", "fund transfer self", "own acc"]
CREDIT_HINT_KEYWORDS = ["salary", "credited", "refund", "cashback", "received", "interest", "dividend", "bonus", "reimbursement"]

EXPENSE_CATEGORY_KEYWORDS = {
    "Food": ["swiggy", "zomato", "restaurant", "dominos", "pizza", "cafe", "food", "dine", "hotel", "mcdonald", "kfc", "starbucks", "haldiram", "eatery", "bakery"],
    "Rent": ["rent", "landlord", "lease"],
    "Transport": ["uber", "ola", "irctc", "petrol", "diesel", "fuel", "metro", "rapido", "fastag", "transport", "railway", "indian oil", "bpcl", "hpcl", "parking", "cab"],
    "Shopping": ["amazon", "flipkart", "myntra", "ajio", "shopping", "mall", "reliance digital", "nykaa", "meesho", "store", "mart"],
    "Entertainment": ["netflix", "hotstar", "spotify", "bookmyshow", "prime video", "movie", "cinema", "pvr", "inox", "sony liv", "zee5", "gaana"],
    "Utilities": ["electricity", "recharge", "airtel", "jio", "vodafone", "broadband", "wifi", "gas bill", "water bill", "dth", "utility", "bses", "mtnl", "bsnl", "postpaid", "prepaid"],
}

INCOME_SOURCE_KEYWORDS = {
    "Salary": ["salary", "sal credit", "payroll"],
    "Interest": ["interest", "int.cr", "int cr", "sb int"],
    "Refund": ["refund", "reversal", "cashback"],
    "Dividend": ["dividend"],
}


def _dedupe_columns(raw_cols):
    """
    Makes a list of column names unique (col, col, col -> col, col_1, col_2).

    Bank exports quite often have duplicate/blank header cells (two empty
    columns, two columns literally both named "Amount", etc). If we leave
    those duplicate names on the dataframe, doing raw_df[some_name] returns
    a mini-DataFrame instead of a single Series wherever that name repeats,
    and assigning that into a single output column later forces pandas to
    reindex it - which throws "Reindexing only valid with uniquely valued
    Index objects" because the duplicate names aren't unique. Deduping the
    names up front avoids that entirely.
    """
    seen = {}
    result = []
    for c in raw_cols:
        c = str(c).strip()
        if c == "" or c.lower() == "nan":
            c = "col"
        if c in seen:
            seen[c] += 1
            result.append(f"{c}_{seen[c]}")
        else:
            seen[c] = 0
            result.append(c)
    return result


def _find_col(columns, synonyms):
    """finds the real column name that best matches one of our synonym lists"""
    for c in columns:
        cl = str(c).strip().lower()
        for s in synonyms:
            if cl == s or (len(s) > 2 and s in cl):
                return c
    return None


def _clean_amount_series(s):
    """turns messy amount strings ('₹1,234.50', '(500.00)', 'Rs. 99') into floats"""
    s = s.astype(str)
    s = s.str.replace(r"[₹,]|Rs\.?", "", regex=True)
    s = s.str.replace(r"\((.*)\)", r"-\1", regex=True)  # (500) means -500
    s = s.str.strip()
    return pd.to_numeric(s, errors="coerce").fillna(0.0)


def _detect_header_row(raw_no_header_df):
    """scans the first ~20 rows of a headerless sheet/table to guess which row is the real header"""
    all_syn = DATE_SYNONYMS + DESC_SYNONYMS + DEBIT_SYNONYMS + CREDIT_SYNONYMS + AMOUNT_SYNONYMS + BALANCE_SYNONYMS + TYPE_SYNONYMS
    best_idx, best_score = None, 0
    for i in range(min(20, len(raw_no_header_df))):
        row_vals = [str(v).strip().lower() for v in raw_no_header_df.iloc[i].tolist()]
        score = sum(1 for v in row_vals for syn in all_syn if v == syn or (len(v) > 2 and syn in v))
        if score > best_score:
            best_score, best_idx = score, i
    return best_idx if best_score >= 2 else None


def normalize_bank_df(raw_df):
    """maps a raw dataframe (any bank's column naming) onto Date/Description/Debit/Credit/Balance"""
    raw_df = raw_df.dropna(axis=1, how="all")
    # defensive dedupe here too, in case a raw_df reaches this function without
    # having gone through one of the _load_* helpers below (e.g. future callers)
    raw_df.columns = _dedupe_columns(raw_df.columns)
    columns = list(raw_df.columns)

    date_col = _find_col(columns, DATE_SYNONYMS)
    desc_col = _find_col(columns, DESC_SYNONYMS)
    debit_col = _find_col(columns, DEBIT_SYNONYMS)
    credit_col = _find_col(columns, CREDIT_SYNONYMS)
    balance_col = _find_col(columns, BALANCE_SYNONYMS)
    amount_col = _find_col(columns, AMOUNT_SYNONYMS)
    type_col = _find_col(columns, TYPE_SYNONYMS)

    if desc_col is None and date_col is None and debit_col is None and credit_col is None and amount_col is None:
        return None  # doesn't look like a transaction table at all

    out = pd.DataFrame()
    out["Date"] = pd.to_datetime(raw_df[date_col], errors="coerce", dayfirst=True) if date_col else pd.NaT
    out["Description"] = raw_df[desc_col].astype(str).str.strip() if desc_col else ""

    if debit_col is not None or credit_col is not None:
        out["Debit"] = _clean_amount_series(raw_df[debit_col]).abs() if debit_col else 0.0
        out["Credit"] = _clean_amount_series(raw_df[credit_col]).abs() if credit_col else 0.0
    elif amount_col is not None:
        amt = _clean_amount_series(raw_df[amount_col])
        if type_col is not None:
            t = raw_df[type_col].astype(str).str.lower()
            is_debit = t.str.contains("dr") | t.str.contains("debit")
            out["Debit"] = amt.abs().where(is_debit, 0.0)
            out["Credit"] = amt.abs().where(~is_debit, 0.0)
        else:
            out["Debit"] = amt.where(amt < 0, 0.0).abs()
            out["Credit"] = amt.where(amt > 0, 0.0)
    else:
        out["Debit"] = 0.0
        out["Credit"] = 0.0

    out["Balance"] = _clean_amount_series(raw_df[balance_col]) if balance_col else pd.NA

    # drop junk rows that have no amount at all (usually blank/footer rows).
    # a row only counts as real data if it has a non-zero debit or credit.
    out = out[(out["Debit"] != 0) | (out["Credit"] != 0)].reset_index(drop=True)
    return out


DATE_REGEX = r"\d{1,2}[\/\-.]\d{1,2}[\/\-.]\d{2,4}|\d{4}[\/\-.]\d{1,2}[\/\-.]\d{1,2}"
AMOUNT_REGEX = r"[-+]?₹?\s?\d[\d,]*\.\d{2}"


def parse_text_fallback(text):
    """
    best-effort line-by-line extraction for unstructured text (scanned PDFs, plain
    txt exports, docx with no real table). looks for a date + trailing amount per line.
    flags every row so the UI can warn the user to double check it.
    """
    rows = []
    for line in text.splitlines():
        line = line.strip()
        if not line:
            continue
        date_match = re.search(DATE_REGEX, line)
        amounts = re.findall(AMOUNT_REGEX, line)
        if not date_match or not amounts:
            continue
        date_str = date_match.group(0)
        amt_str = amounts[-1]
        desc = line.replace(date_str, "")
        for a in amounts:
            desc = desc.replace(a, "")
        desc = re.sub(r"\s+", " ", desc).strip(" -|,:")
        amt_val = float(re.sub(r"[₹,\s]", "", amt_str))
        if amt_val <= 0:
            continue
        rows.append({"Date": date_str, "Description": desc if desc else "(unlabelled transaction)", "Amount": amt_val})

    if not rows:
        return None

    df = pd.DataFrame(rows)
    df["Date"] = pd.to_datetime(df["Date"], errors="coerce", dayfirst=True)

    # no reliable debit/credit signal in free text - guess Credit if the description
    # itself hints at income (salary/refund/etc), else assume Debit (most personal
    # statement lines are spends). the review table lets the user flip this per row.
    desc_l = df["Description"].str.lower()
    is_credit_guess = desc_l.apply(lambda d: any(k in d for k in CREDIT_HINT_KEYWORDS))
    out = pd.DataFrame()
    out["Date"] = df["Date"]
    out["Description"] = df["Description"]
    out["Debit"] = df["Amount"].where(~is_credit_guess, 0.0)
    out["Credit"] = df["Amount"].where(is_credit_guess, 0.0)
    out["Balance"] = pd.NA
    return out


def _load_csv(uploaded_file):
    """
    Reads a bank-exported CSV.

    IMPORTANT: we deliberately try the header-agnostic parser FIRST, not as a
    fallback. Real bank CSV exports usually have junk rows above the real
    header (account title, "A/C No: xxxx", blank lines, etc). A plain
    pd.read_csv() call happily "succeeds" on these files - it just uses the
    junk row as the header - so it never raises an exception and the smart
    detection logic used to never even run. That's what was breaking imports:
    parse_bank_file() would get a raw_df back with garbage column names,
    normalize_bank_df() couldn't find Date/Debit/Credit/etc among them and
    returned None, and there was no fallback_text either (because _load_csv
    "succeeded"), so parse_bank_file() had nothing to fall back to and
    reported "Couldn't find any transactions."
    """
    raw_bytes = uploaded_file.getvalue()
    text = None
    for enc in ("utf-8-sig", "utf-8", "latin-1"):
        try:
            text = raw_bytes.decode(enc)
            break
        except Exception:
            continue
    if text is None:
        text = raw_bytes.decode("utf-8", errors="replace")

    # 1) try robust, header-agnostic parsing FIRST - handles both clean
    # exports and ones with junk rows above the real header (most Indian
    # bank CSV exports look like this)
    try:
        import csv as csv_module
        reader = csv_module.reader(io.StringIO(text))
        rows = [row for row in reader if any(cell.strip() for cell in row)]
        if rows:
            max_cols = max(len(r) for r in rows)
            rows = [r + [""] * (max_cols - len(r)) for r in rows]
            no_header_df = pd.DataFrame(rows)
            header_row = _detect_header_row(no_header_df)
            if header_row is not None:
                header = no_header_df.iloc[header_row].tolist()
                data = no_header_df.iloc[header_row + 1:].reset_index(drop=True)
                data.columns = _dedupe_columns(header)
                return data, None
    except Exception:
        pass

    # 2) fallback: plain pandas read, for clean CSVs that for whatever reason
    # didn't score high enough on the synonym check above
    try:
        raw_df = pd.read_csv(io.StringIO(text))
        if raw_df.shape[1] >= 2:
            raw_df.columns = _dedupe_columns(raw_df.columns)
            return raw_df, None
    except Exception:
        pass

    # 3) nothing structured worked - hand back raw text for the free-text parser
    return None, text


def _load_excel(uploaded_file):
    xls = pd.ExcelFile(uploaded_file)
    for sheet in xls.sheet_names:
        no_header = pd.read_excel(xls, sheet_name=sheet, header=None)
        header_row = _detect_header_row(no_header)
        if header_row is not None:
            df = pd.read_excel(xls, sheet_name=sheet, header=header_row)
            if len(df) > 0:
                df.columns = _dedupe_columns(df.columns)
                return df, None
    # nothing scored well enough - just use the first sheet as-is
    df = pd.read_excel(xls, sheet_name=xls.sheet_names[0])
    df.columns = _dedupe_columns(df.columns)
    return df, None


def _load_pdf(uploaded_file):
    import pdfplumber
    table_frames, text_parts = [], []
    with pdfplumber.open(uploaded_file) as pdf:
        for page in pdf.pages:
            try:
                for t in page.extract_tables():
                    if t and len(t) >= 2:
                        table_frames.append(pd.DataFrame(t[1:], columns=_dedupe_columns(t[0])))
            except Exception:
                pass
            try:
                text_parts.append(page.extract_text() or "")
            except Exception:
                pass
    if table_frames:
        combined = pd.concat(table_frames, ignore_index=True, sort=False)
        combined.columns = _dedupe_columns(combined.columns)
        return combined, None
    return None, "\n".join(text_parts)


def _load_docx(uploaded_file):
    import docx
    document = docx.Document(uploaded_file)
    table_frames = []
    for table in document.tables:
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        if len(rows) >= 2:
            table_frames.append(pd.DataFrame(rows[1:], columns=_dedupe_columns(rows[0])))
    if table_frames:
        combined = pd.concat(table_frames, ignore_index=True, sort=False)
        combined.columns = _dedupe_columns(combined.columns)
        return combined, None
    text = "\n".join(p.text for p in document.paragraphs)
    return None, text


def _load_xml(uploaded_file):
    try:
        df = pd.read_xml(uploaded_file)
        df.columns = _dedupe_columns(df.columns)
        return df, None
    except Exception:
        pass
    try:
        import xml.etree.ElementTree as ET
        raw_bytes = uploaded_file.getvalue()
        root = ET.fromstring(raw_bytes)
        rows = []
        for child in root.iter():
            grandchildren = list(child)
            if grandchildren and all(len(list(gc)) == 0 for gc in grandchildren):
                row = {gc.tag: (gc.text or "").strip() for gc in grandchildren}
                if row:
                    rows.append(row)
        if rows:
            return pd.DataFrame(rows), None
        return None, raw_bytes.decode("utf-8", errors="replace")
    except Exception:
        raw_bytes = uploaded_file.getvalue()
        return None, raw_bytes.decode("utf-8", errors="replace")


def classify_type(description, debit, credit, ambiguous=False):
    desc_l = str(description).lower()
    if any(k in desc_l for k in SELF_TRANSFER_KEYWORDS):
        return "Self Transfer"
    if ambiguous:
        return "Income" if any(k in desc_l for k in CREDIT_HINT_KEYWORDS) else "Expense"
    if credit > 0 and debit == 0:
        return "Income"
    if debit > 0 and credit == 0:
        return "Expense"
    return "Income" if credit >= debit else "Expense"


def classify_category(description, type_):
    desc_l = str(description).lower()
    if type_ == "Self Transfer":
        return "Self Transfer"
    if type_ == "Income":
        for src, kws in INCOME_SOURCE_KEYWORDS.items():
            if any(k in desc_l for k in kws):
                return src
        return "Other Income"
    for cat, kws in EXPENSE_CATEGORY_KEYWORDS.items():
        if any(k in desc_l for k in kws):
            return cat
    return "Other"


def parse_bank_file(uploaded_file):
    """
    main entry point: reads any supported bank statement file and returns
    (review_df, mode) where review_df has Date/Description/Amount/Type/Category
    and mode is 'table' (clean structured parse) or 'text' (best-effort fallback)
    """
    name = uploaded_file.name.lower()
    ext = name.rsplit(".", 1)[-1] if "." in name else ""

    raw_df, fallback_text = None, None

    if ext == "csv":
        raw_df, fallback_text = _load_csv(uploaded_file)
    elif ext in ("xlsx", "xls"):
        raw_df, fallback_text = _load_excel(uploaded_file)
    elif ext == "pdf":
        raw_df, fallback_text = _load_pdf(uploaded_file)
    elif ext == "docx":
        raw_df, fallback_text = _load_docx(uploaded_file)
    elif ext == "xml":
        raw_df, fallback_text = _load_xml(uploaded_file)
    elif ext == "txt":
        fallback_text = uploaded_file.getvalue().decode("utf-8", errors="replace")
    else:
        raise ValueError(f"unsupported file type: .{ext}")

    norm, mode = None, "table"
    if raw_df is not None:
        norm = normalize_bank_df(raw_df)

    if (norm is None or norm.empty) and fallback_text:
        norm = parse_text_fallback(fallback_text)
        mode = "text"

    if norm is None or norm.empty:
        return pd.DataFrame(columns=["Date", "Description", "Amount", "Type", "Category"]), "empty"

    ambiguous = (mode == "text")
    review_rows = []
    for _, r in norm.iterrows():
        debit, credit = float(r["Debit"]), float(r["Credit"])
        t = classify_type(r["Description"], debit, credit, ambiguous=ambiguous)
        c = classify_category(r["Description"], t)
        review_rows.append({
            "Date": r["Date"],
            "Description": r["Description"] if str(r["Description"]).strip() else "(no description)",
            "Amount": debit if debit > 0 else credit,
            "Type": t,
            "Category": c,
        })
    review_df = pd.DataFrame(review_rows)
    return review_df, mode


# =========================================================
# SESSION STATE
# =========================================================
if "expenses" not in st.session_state:
    st.session_state.expenses = load_expenses()

if "chat_history" not in st.session_state:
    st.session_state.chat_history = []

if "budgets" not in st.session_state:
    st.session_state.budgets = {}

if "scanned_items" not in st.session_state:
    st.session_state.scanned_items = pd.DataFrame(columns=cols)

if "income_entries" not in st.session_state:
    st.session_state.income_entries = pd.DataFrame(columns=["Date", "Source", "Remark", "Amount"])

if "transfers" not in st.session_state:
    st.session_state.transfers = pd.DataFrame(columns=["Date", "Description", "Amount"])

if "bank_import_preview" not in st.session_state:
    st.session_state.bank_import_preview = None

if "bank_import_mode" not in st.session_state:
    st.session_state.bank_import_mode = None

# =========================================================
# SIDEBAR
# =========================================================
with st.sidebar:
    st.markdown("### 💸 Budget Planner")
    st.caption("Track spends, scan bills, get AI tips")
    st.divider()

    st.markdown("**API Key**")
    # checks Streamlit Cloud secrets first (st.secrets), then falls back to .env / local env var
    try:
        env_key = st.secrets["GEMINI_API_KEY"]
    except (KeyError, FileNotFoundError):
        env_key = os.getenv("GEMINI_API_KEY")
    GEMINI_API_KEY = st.text_input("GEMINI_API_KEY", type="password", value=env_key if env_key else "", label_visibility="collapsed")

    if GEMINI_API_KEY:
        st.success("API key loaded ✅")
    else:
        st.info("Paste your Gemini API key to unlock AI features")
        st.caption("Get one free at [aistudio.google.com](https://aistudio.google.com/app/apikey)")

    st.divider()
    st.markdown("**Quick Stats**")
    _total_exp_sb = st.session_state.expenses["Amount"].sum() if not st.session_state.expenses.empty else 0.0
    _total_inc_imported_sb = st.session_state.income_entries["Amount"].sum() if not st.session_state.income_entries.empty else 0.0
    st.markdown(
        f'<div class="side-stat"><span>Total Expense</span><b>{money(_total_exp_sb)}</b></div>'
        f'<div class="side-stat"><span>Entries</span><b>{len(st.session_state.expenses)}</b></div>'
        f'<div class="side-stat"><span>Imported Income</span><b>{money(_total_inc_imported_sb)}</b></div>',
        unsafe_allow_html=True
    )

    st.divider()
    st.caption("Made with Streamlit + Gemini ✨")

# =========================================================
# HERO HEADER
# =========================================================
st.markdown("""
<div class="hero">
    <h1>AI Personal Budget Planner 💰</h1>
    <p>Track your income and expenses, scan bills automatically, and get AI-powered budgeting advice.</p>
</div>
""", unsafe_allow_html=True)

# ---------------- INCOME ----------------
with st.container(border=True):
    st.markdown('<div class="section-title">💵 Monthly Income</div>', unsafe_allow_html=True)
    base_income = st.number_input("Enter Monthly Income (₹)", min_value=0.0, step=500.0, label_visibility="collapsed")
    _imported_income = st.session_state.income_entries["Amount"].sum() if not st.session_state.income_entries.empty else 0.0
    if _imported_income > 0:
        st.caption(f"+ {money(_imported_income)} from imported bank statement income → total counted below")

# income used everywhere else in the app = manual entry + whatever got imported
# from bank statements (salary/interest/refund rows etc)
income = base_income + _imported_income

# persistent summary, visible no matter which tab is open
_total_exp_top = st.session_state.expenses["Amount"].sum() if not st.session_state.expenses.empty else 0.0
render_summary_bar(income, _total_exp_top, income - _total_exp_top, income - _total_exp_top)

# using tabs so everything doesnt look like one giant scrolling page
tab_add, tab_import, tab_dash, tab_forecast, tab_advisor, tab_chat, tab_share = st.tabs(
    ["➕ Add Expense", "📥 Import Statement", "📊 Dashboard", "📈 Forecast", "🤖 AI Advisor", "💬 Chat", "📤 Share/Export"]
)

# =========================================================
# TAB 1 - ADD EXPENSE
# =========================================================
with tab_add:
    with st.container(border=True):
        st.markdown('<div class="section-title">✍️ Add Manually</div>', unsafe_allow_html=True)

        c1, c2, c3, c4 = st.columns([1.2, 1.4, 1, 0.8])

        with c1:
            category = st.selectbox("Category", CATEGORY_OPTIONS)
            custom_cat = ""
            if category == "Other":
                custom_cat = st.text_input("Custom Category Name")
        with c2:
            remark = st.text_input("Remark (optional)")
        with c3:
            amount = st.number_input("Amount (₹)", min_value=0.0, step=50.0)
        with c4:
            st.write("")
            st.write("")
            add_btn = st.button("Add ➕", use_container_width=True)

        if add_btn:
            final_cat = custom_cat.strip() if category == "Other" and custom_cat.strip() != "" else category

            if amount > 0:
                new_row = pd.DataFrame([{"Category": final_cat, "Remark": remark, "Amount": amount, "Date": pd.Timestamp.now().normalize()}])
                st.session_state.expenses = pd.concat([st.session_state.expenses, new_row], ignore_index=True)
                st.session_state.expenses["Amount"] = pd.to_numeric(st.session_state.expenses["Amount"], errors="coerce").fillna(0.0)
                save_expenses(st.session_state.expenses)
                st.toast(f"Added {final_cat} — {money(amount)}", icon="✅")
            else:
                st.warning("amount should be more than 0")

    with st.container(border=True):
        st.markdown('<div class="section-title">📷 Scan a Bill / Receipt</div>', unsafe_allow_html=True)
        st.markdown('<div class="section-caption">Upload or capture a bill photo — AI reads it and fills in the expenses for you.</div>', unsafe_allow_html=True)

        if not GEMINI_API_KEY:
            st.info("Add your Gemini API key in the sidebar to use bill scanning")
        else:
            scan_c1, scan_c2 = st.columns(2)
            with scan_c1:
                uploaded_bill = st.file_uploader("Upload bill image", type=["jpg", "jpeg", "png", "webp"], key="bill_upload")
            with scan_c2:
                camera_bill = st.camera_input("Or take a photo")

            bill_file = camera_bill if camera_bill is not None else uploaded_bill

            if bill_file is not None:
                pc1, pc2 = st.columns([1, 2])
                with pc1:
                    st.image(bill_file, caption="Preview", use_container_width=True)
                with pc2:
                    st.write("")
                    if st.button("🔍 Scan Bill with AI", use_container_width=True):
                        with st.spinner("Reading your bill..."):
                            try:
                                img_bytes = bill_file.getvalue()
                                mime = bill_file.type if getattr(bill_file, "type", None) else "image/jpeg"
                                extracted = extract_expenses_from_bill(img_bytes, mime)

                                if not extracted:
                                    st.warning("Couldn't read any expense lines from this image — try a clearer photo")
                                else:
                                    scanned_df = pd.DataFrame(extracted)
                                    scanned_df["Date"] = pd.Timestamp.now().normalize()
                                    st.session_state.scanned_items = scanned_df
                                    st.success(f"Found {len(extracted)} item(s) — review below before adding")
                            except json.JSONDecodeError:
                                st.error("AI didn't return clean data, try scanning again or use a clearer photo")
                            except Exception as e:
                                st.error("Error: " + str(e))

            if not st.session_state.scanned_items.empty:
                st.write("**Review scanned items** (edit if needed):")
                reviewed = st.data_editor(
                    st.session_state.scanned_items,
                    use_container_width=True,
                    num_rows="dynamic",
                    key="scanned_editor",
                )

                rc1, rc2 = st.columns(2)
                with rc1:
                    if st.button("✅ Add Scanned Items to Expenses", use_container_width=True):
                        reviewed["Amount"] = pd.to_numeric(reviewed["Amount"], errors="coerce").fillna(0.0)
                        reviewed = reviewed[reviewed["Amount"] > 0]
                        st.session_state.expenses = pd.concat([st.session_state.expenses, reviewed], ignore_index=True)
                        save_expenses(st.session_state.expenses)
                        st.session_state.scanned_items = pd.DataFrame(columns=cols)
                        st.toast("Scanned items added!", icon="🎉")
                        st.rerun()
                with rc2:
                    if st.button("❌ Discard Scanned Items", use_container_width=True):
                        st.session_state.scanned_items = pd.DataFrame(columns=cols)
                        st.rerun()

    with st.container(border=True):
        st.markdown('<div class="section-title">📋 Expense Table</div>', unsafe_allow_html=True)

        if st.session_state.expenses.empty:
            empty_state("🧾", "No expenses yet", "Add one manually or scan a bill above to get started")
        else:
            edited = st.data_editor(st.session_state.expenses, use_container_width=True, num_rows="dynamic")

            col_save, col_clear = st.columns(2)
            with col_save:
                if st.button("💾 Save Changes", use_container_width=True):
                    edited["Amount"] = pd.to_numeric(edited["Amount"], errors="coerce").fillna(0.0)
                    st.session_state.expenses = edited
                    save_expenses(st.session_state.expenses)
                    st.toast("Saved!", icon="💾")
            with col_clear:
                if st.button("🗑️ Clear All", use_container_width=True):
                    st.session_state.expenses = pd.DataFrame(columns=cols)
                    save_expenses(st.session_state.expenses)
                    st.rerun()

    with st.container(border=True):
        st.markdown('<div class="section-title">🎯 Category Budget Limits</div>', unsafe_allow_html=True)
        st.markdown('<div class="section-caption">Optional — resets when you close the app (not saved to file yet)</div>', unsafe_allow_html=True)

        bcol = st.columns(4)
        limit_cats = ["Food", "Rent", "Transport", "Shopping"]
        for i in range(len(limit_cats)):
            cat = limit_cats[i]
            with bcol[i]:
                st.session_state.budgets[cat] = st.number_input(cat + " Limit", min_value=0.0, step=500.0, key="lim_" + cat)

# =========================================================
# TAB - IMPORT BANK STATEMENT
# =========================================================
with tab_import:
    with st.container(border=True):
        st.markdown('<div class="section-title">📥 Upload Bank Statement</div>', unsafe_allow_html=True)
        st.markdown('<div class="section-caption">CSV, Excel, PDF, Word (.docx), XML or plain text - we auto-detect income, expenses & self-transfers</div>', unsafe_allow_html=True)

        uploaded_stmt = st.file_uploader(
            "Upload statement", type=["csv", "xlsx", "xls", "pdf", "docx", "xml", "txt"],
            key="bank_stmt_upload", label_visibility="collapsed"
        )

        if uploaded_stmt is not None and st.button("🔍 Analyze Statement", use_container_width=True):
            with st.spinner("Reading & analyzing your statement..."):
                try:
                    review_df, mode = parse_bank_file(uploaded_stmt)
                    if review_df.empty:
                        st.error("Couldn't find any transactions in this file. A CSV or Excel export from your bank usually parses most reliably.")
                        st.session_state.bank_import_preview = None
                    else:
                        st.session_state.bank_import_preview = review_df
                        st.session_state.bank_import_mode = mode
                        st.success(f"Found {len(review_df)} transaction(s)")
                        if mode == "text":
                            st.warning("Couldn't detect a clean table in this file — used best-effort text extraction. Please double-check Type / Category / Amount below before importing.")
                except ValueError as e:
                    st.error(str(e))
                except ImportError:
                    st.error("A required library for this file type isn't installed. Make sure requirements.txt includes pdfplumber / python-docx / openpyxl / lxml.")
                except Exception as e:
                    st.error(f"Couldn't read this file: {e}")

    preview = st.session_state.bank_import_preview
    if preview is not None and not preview.empty:
        # ---------------- QUICK EDA ----------------
        with st.container(border=True):
            st.markdown('<div class="section-title">🔎 Quick EDA</div>', unsafe_allow_html=True)

            total_credit = preview.loc[preview["Type"] == "Income", "Amount"].sum()
            total_debit = preview.loc[preview["Type"] == "Expense", "Amount"].sum()
            n_self = int((preview["Type"] == "Self Transfer").sum())
            valid_dates = pd.to_datetime(preview["Date"], errors="coerce").dropna()
            date_range_str = f"{valid_dates.min().date()} → {valid_dates.max().date()}" if len(valid_dates) else "unknown"

            e1, e2, e3, e4 = st.columns(4)
            e1.markdown(metric_card_html("Transactions", str(len(preview)), date_range_str, "#4F46E5"), unsafe_allow_html=True)
            e2.markdown(metric_card_html("Total Credit", money(total_credit), "detected income", "#22C55E"), unsafe_allow_html=True)
            e3.markdown(metric_card_html("Total Debit", money(total_debit), "detected expense", "#EF4444"), unsafe_allow_html=True)
            e4.markdown(metric_card_html("Self Transfers", str(n_self), "excluded from totals", "#64748B"), unsafe_allow_html=True)
            st.markdown("<br>", unsafe_allow_html=True)

            if len(valid_dates) >= 2:
                tmp = preview.dropna(subset=["Date"]).copy()
                tmp["Date"] = pd.to_datetime(tmp["Date"], errors="coerce")
                tmp = tmp.dropna(subset=["Date"])
                tmp["Month"] = tmp["Date"].dt.to_period("M").astype(str)
                monthly = tmp[tmp["Type"] != "Self Transfer"].groupby(["Month", "Type"])["Amount"].sum().reset_index()
                if not monthly.empty:
                    fig_m = px.bar(monthly, x="Month", y="Amount", color="Type", barmode="group",
                                    color_discrete_map={"Income": "#22C55E", "Expense": "#EF4444"})
                    fig_m.update_layout(margin=dict(t=10, b=10, l=10, r=10), height=300, plot_bgcolor="white")
                    st.plotly_chart(fig_m, use_container_width=True)

            exp_only = preview[preview["Type"] == "Expense"]
            if not exp_only.empty:
                cat_prev = exp_only.groupby("Category")["Amount"].sum().reset_index().sort_values("Amount", ascending=False)
                fig_c = px.bar(cat_prev, x="Amount", y="Category", orientation="h", color="Category",
                                color_discrete_map=CATEGORY_COLORS)
                fig_c.update_layout(showlegend=False, margin=dict(t=10, b=10, l=10, r=10), height=280,
                                     xaxis_title=None, yaxis_title=None, plot_bgcolor="white")
                st.plotly_chart(fig_c, use_container_width=True)

        # ---------------- REVIEW & EDIT ----------------
        with st.container(border=True):
            st.markdown('<div class="section-title">✏️ Review & Correct</div>', unsafe_allow_html=True)
            st.markdown('<div class="section-caption">We auto-detected Type and Category from each description — fix anything that looks off before importing</div>', unsafe_allow_html=True)

            all_category_options = sorted(set(CATEGORY_OPTIONS + INCOME_SOURCE_OPTIONS + ["Self Transfer"]))
            edited_preview = st.data_editor(
                preview,
                use_container_width=True,
                num_rows="dynamic",
                column_config={
                    "Type": st.column_config.SelectboxColumn("Type", options=["Income", "Expense", "Self Transfer"]),
                    "Category": st.column_config.SelectboxColumn("Category", options=all_category_options),
                    "Amount": st.column_config.NumberColumn("Amount", min_value=0.0, format="₹%.2f"),
                    "Date": st.column_config.DateColumn("Date"),
                },
                key="bank_import_editor",
            )

            ic1, ic2 = st.columns(2)
            with ic1:
                if st.button("✅ Import These Transactions", use_container_width=True):
                    edited = edited_preview.copy()
                    edited["Amount"] = pd.to_numeric(edited["Amount"], errors="coerce").fillna(0.0)
                    edited = edited[edited["Amount"] > 0]

                    exp_rows = edited[edited["Type"] == "Expense"]
                    inc_rows = edited[edited["Type"] == "Income"]
                    self_rows = edited[edited["Type"] == "Self Transfer"]

                    if not exp_rows.empty:
                        new_exp = pd.DataFrame({
                            "Category": exp_rows["Category"],
                            "Remark": exp_rows["Description"],
                            "Amount": exp_rows["Amount"],
                            "Date": exp_rows["Date"],
                        })
                        st.session_state.expenses = pd.concat([st.session_state.expenses, new_exp], ignore_index=True)
                        st.session_state.expenses["Amount"] = pd.to_numeric(st.session_state.expenses["Amount"], errors="coerce").fillna(0.0)
                        save_expenses(st.session_state.expenses)

                    if not inc_rows.empty:
                        new_inc = pd.DataFrame({
                            "Date": inc_rows["Date"],
                            "Source": inc_rows["Category"],
                            "Remark": inc_rows["Description"],
                            "Amount": inc_rows["Amount"],
                        })
                        st.session_state.income_entries = pd.concat([st.session_state.income_entries, new_inc], ignore_index=True)

                    if not self_rows.empty:
                        new_self = pd.DataFrame({
                            "Date": self_rows["Date"],
                            "Description": self_rows["Description"],
                            "Amount": self_rows["Amount"],
                        })
                        st.session_state.transfers = pd.concat([st.session_state.transfers, new_self], ignore_index=True)

                    st.session_state.bank_import_preview = None
                    st.toast(f"Imported {len(exp_rows)} expense(s), {len(inc_rows)} income(s), {len(self_rows)} self-transfer(s)", icon="🎉")
                    st.rerun()
            with ic2:
                if st.button("❌ Discard", use_container_width=True):
                    st.session_state.bank_import_preview = None
                    st.rerun()
    else:
        empty_state("📥", "Upload a bank statement", "CSV, Excel, PDF, Word or XML — we'll auto-detect income, expenses & self-transfers")

    if not st.session_state.income_entries.empty:
        with st.expander(f"💰 {len(st.session_state.income_entries)} imported income entry/entries (added to Monthly Income above)"):
            st.dataframe(st.session_state.income_entries, use_container_width=True)
            if st.button("🗑️ Clear Imported Income", key="clear_income_entries"):
                st.session_state.income_entries = pd.DataFrame(columns=["Date", "Source", "Remark", "Amount"])
                st.rerun()

    if not st.session_state.transfers.empty:
        with st.expander(f"🔁 {len(st.session_state.transfers)} self-transfer(s) logged (excluded from income/expense totals)"):
            st.dataframe(st.session_state.transfers, use_container_width=True)
            if st.button("🗑️ Clear Transfer Log", key="clear_transfers"):
                st.session_state.transfers = pd.DataFrame(columns=["Date", "Description", "Amount"])
                st.rerun()

# =========================================================
# TAB 2 - DASHBOARD
# =========================================================
with tab_dash:
    total_exp = st.session_state.expenses["Amount"].sum() if not st.session_state.expenses.empty else 0.0
    balance = income - total_exp
    savings = balance

    if not st.session_state.expenses.empty:
        cat_totals = st.session_state.expenses.groupby("Category")["Amount"].sum().reset_index()
        colors = [CATEGORY_COLORS.get(c, "#94A3B8") for c in cat_totals["Category"]]

        colp, colb = st.columns(2)
        with colp:
            with st.container(border=True):
                st.markdown('<div class="section-title">🥧 Expense Breakdown</div>', unsafe_allow_html=True)
                fig1 = px.pie(cat_totals, names="Category", values="Amount", hole=0.55,
                              color="Category", color_discrete_map=CATEGORY_COLORS)
                fig1.update_traces(textinfo="percent+label", textfont_size=12)
                fig1.update_layout(showlegend=False, margin=dict(t=10, b=10, l=10, r=10), height=320)
                st.plotly_chart(fig1, use_container_width=True)

        with colb:
            with st.container(border=True):
                st.markdown('<div class="section-title">📊 Spend by Category</div>', unsafe_allow_html=True)
                fig2 = px.bar(cat_totals.sort_values("Amount"), x="Amount", y="Category", orientation="h",
                              color="Category", color_discrete_map=CATEGORY_COLORS, text="Amount")
                fig2.update_traces(texttemplate="₹%{text:,.0f}", textposition="outside")
                fig2.update_layout(showlegend=False, margin=dict(t=10, b=10, l=10, r=10), height=320,
                                    xaxis_title=None, yaxis_title=None)
                st.plotly_chart(fig2, use_container_width=True)

        if st.session_state.budgets and any(v > 0 for v in st.session_state.budgets.values()):
            with st.container(border=True):
                st.markdown('<div class="section-title">🎯 Budget Limit Check</div>', unsafe_allow_html=True)
                for cat in st.session_state.budgets:
                    lim = st.session_state.budgets[cat]
                    if lim > 0:
                        spent = st.session_state.expenses.loc[st.session_state.expenses["Category"] == cat, "Amount"].sum()
                        budget_bar(cat, spent, lim)

        _dated = st.session_state.expenses.dropna(subset=["Date"]) if "Date" in st.session_state.expenses.columns else pd.DataFrame()
        if len(_dated) >= 2:
            with st.container(border=True):
                st.markdown('<div class="section-title">📅 Spending Trend</div>', unsafe_allow_html=True)
                st.markdown('<div class="section-caption">By day, across manually added, scanned & imported expenses</div>', unsafe_allow_html=True)
                trend = _dated.copy()
                trend["Date"] = pd.to_datetime(trend["Date"], errors="coerce")
                trend = trend.dropna(subset=["Date"]).groupby(trend["Date"].dt.date)["Amount"].sum().reset_index()
                fig_t = px.line(trend, x="Date", y="Amount", markers=True)
                fig_t.update_traces(line=dict(color="#4F46E5", width=3), marker=dict(size=6, color="#4F46E5"))
                fig_t.update_layout(margin=dict(t=10, b=10, l=10, r=10), height=280, plot_bgcolor="white", yaxis_title="₹")
                st.plotly_chart(fig_t, use_container_width=True)
    else:
        with st.container(border=True):
            empty_state("📊", "Nothing to chart yet", "Add some expenses in the Add Expense tab first")

    with st.container(border=True):
        st.markdown('<div class="section-title">💡 Saving Tips</div>', unsafe_allow_html=True)
        if income > 0:
            rate = (savings / income) * 100
            if rate < 20:
                st.warning(f"Savings is only {rate:.1f}% of income — try to save at least 20%")
            else:
                st.success(f"Nice! You're saving {rate:.1f}% of income")

            if not st.session_state.expenses.empty:
                food_amt = st.session_state.expenses.loc[st.session_state.expenses["Category"] == "Food", "Amount"].sum()
                if food_amt > 0.15 * income:
                    st.warning("Food expense is quite high — try cooking at home more often")

            if total_exp > income:
                st.error("You are spending more than you earn this month!")

            st.markdown(
"""- Try to save money as soon as salary comes in
- Track small expenses too — they add up
- Cancel subscriptions you don't use
- Keep some emergency fund"""
            )
        else:
            st.info("Enter income above to see personalized tips")

# =========================================================
# TAB 3 - FORECAST
# =========================================================
with tab_forecast:
    total_exp = st.session_state.expenses["Amount"].sum() if not st.session_state.expenses.empty else 0.0

    with st.container(border=True):
        st.markdown('<div class="section-title">📈 Forecast</div>', unsafe_allow_html=True)
        st.markdown('<div class="section-caption">Simple projection based on your current numbers</div>', unsafe_allow_html=True)

        if not st.session_state.expenses.empty and income > 0:
            horizon = st.radio("Forecast for", ["Monthly", "Yearly"], horizontal=True, label_visibility="collapsed")
            proj_saving = income - total_exp

            if horizon == "Monthly":
                x = list(range(1, 7))
                y = [proj_saving * i for i in x]
                xlabel = "Months ahead"
            else:
                x = list(range(1, 6))
                y = [proj_saving * 12 * i for i in x]
                xlabel = "Years ahead"

            fig4 = go.Figure()
            fig4.add_trace(go.Scatter(
                x=x, y=y, mode="lines+markers", line=dict(color="#22C55E", width=3),
                marker=dict(size=8, color="#22C55E"), fill="tozeroy",
                fillcolor="rgba(34,197,94,0.12)"
            ))
            fig4.update_layout(
                xaxis_title=xlabel, yaxis_title="Projected Savings (₹)",
                margin=dict(t=20, b=10, l=10, r=10), height=340,
                plot_bgcolor="white"
            )
            st.plotly_chart(fig4, use_container_width=True)

            f1, f2, f3 = st.columns(3)
            f1.markdown(metric_card_html("Current Monthly Expense", money(total_exp), "", "#F97316"), unsafe_allow_html=True)
            f2.markdown(metric_card_html("Projected Yearly Expense", money(total_exp * 12), "", "#EF4444"), unsafe_allow_html=True)
            f3.markdown(metric_card_html("Projected Yearly Savings", money(proj_saving * 12), "", "#22C55E"), unsafe_allow_html=True)
        else:
            empty_state("📈", "Nothing to forecast yet", "Enter income and add expenses first")

# =========================================================
# TAB 4 - AI ADVISOR
# =========================================================
with tab_advisor:
    total_exp = st.session_state.expenses["Amount"].sum() if not st.session_state.expenses.empty else 0.0
    balance = income - total_exp

    def get_context():
        exp_summary = st.session_state.expenses.groupby("Category")["Amount"].sum().to_string() if not st.session_state.expenses.empty else "No expenses recorded"
        return f"""
        Monthly Income: {income}
        Total Expense: {total_exp}
        Balance: {balance}
        Expense Breakdown:
        {exp_summary}
        """

    with st.container(border=True):
        st.markdown('<div class="section-title">🤖 AI Financial Advisor</div>', unsafe_allow_html=True)

        if st.button("✨ Get AI Financial Advice", use_container_width=True):
            if not GEMINI_API_KEY:
                st.error("Give API key in sidebar first")
            elif income == 0:
                st.warning("Enter monthly income first")
            else:
                with st.spinner("Asking Gemini..."):
                    try:
                        advisor_prompt = ChatPromptTemplate.from_template(
                            "You are a friendly personal finance advisor. {context} "
                            "Give me 3-5 short practical tips to improve my budget and savings."
                        )
                        chain = advisor_prompt | get_gemini_model() | StrOutputParser()
                        result = chain.invoke({"context": get_context()})
                        with st.chat_message("assistant"):
                            st.markdown(result)
                    except Exception as e:
                        st.error("Error: " + str(e))

    st.warning("⚠️ Disclaimer: Yeh AI advice sirf general guidance ke liye hai. Kripya koi bhi bada financial decision lene se pehle ek certified financial adviser se consult karo.")

# =========================================================
# TAB 5 - CHAT
# =========================================================
with tab_chat:
    total_exp = st.session_state.expenses["Amount"].sum() if not st.session_state.expenses.empty else 0.0
    balance = income - total_exp

    with st.container(border=True):
        st.markdown('<div class="section-title">💬 Chat with your Planner</div>', unsafe_allow_html=True)
        st.markdown('<div class="section-caption">Ask stuff like "where can I cut cost" or "can I afford a ₹5000 trip"</div>', unsafe_allow_html=True)

        for m in st.session_state.chat_history:
            with st.chat_message(m["role"]):
                st.write(m["content"])

        q = st.chat_input("Ask something...")

        if q:
            st.session_state.chat_history.append({"role": "user", "content": q})
            with st.chat_message("user"):
                st.write(q)

            if not GEMINI_API_KEY:
                ans = "Give API key in sidebar to use chat"
            else:
                try:
                    exp_summary = st.session_state.expenses.groupby("Category")["Amount"].sum().to_string() if not st.session_state.expenses.empty else "No expenses recorded"
                    chat_prompt = ChatPromptTemplate.from_template(
                        """You are a budget planning assistant.
                        Income: {income}, Expense: {total_exp}, Balance: {balance}
                        Expense breakdown: {exp_summary}
                        Question: {question}
                        Answer short and clear."""
                    )
                    chain = chat_prompt | get_gemini_model() | StrOutputParser()
                    ans = chain.invoke({
                        "income": income,
                        "total_exp": total_exp,
                        "balance": balance,
                        "exp_summary": exp_summary,
                        "question": q
                    })
                except Exception as e:
                    ans = "Error: " + str(e)

            st.session_state.chat_history.append({"role": "assistant", "content": ans})
            with st.chat_message("assistant"):
                st.write(ans)

# =========================================================
# TAB 6 - SHARE / EXPORT
# =========================================================
with tab_share:
    total_exp = st.session_state.expenses["Amount"].sum() if not st.session_state.expenses.empty else 0.0
    balance = income - total_exp

    with st.container(border=True):
        st.markdown('<div class="section-title">⬇️ Export Data</div>', unsafe_allow_html=True)
        if not st.session_state.expenses.empty:
            csv_data = st.session_state.expenses.to_csv(index=False)
            st.download_button("Download CSV", data=csv_data, file_name="my_expenses.csv", mime="text/csv", use_container_width=True)
        else:
            empty_state("⬇️", "Nothing to export yet", "Add expenses first to download")

    exp_summary = st.session_state.expenses.groupby("Category")["Amount"].sum().to_string() if not st.session_state.expenses.empty else "No expenses recorded"
    report_text = f"""AI Personal Budget Planner Report

Monthly Income: Rs.{income:,.2f}
Total Expense: Rs.{total_exp:,.2f}
Balance: Rs.{balance:,.2f}

Expense Breakdown:
{exp_summary}

- generated by AI Personal Budget Planner"""

    colA, colB = st.columns(2)

    with colA:
        with st.container(border=True):
            st.markdown('<div class="section-title">📧 Send via Email</div>', unsafe_allow_html=True)
            to_email = st.text_input("Recipient Email")

            subject = "Your Monthly Budget Report"
            mail_subject = urllib.parse.quote(subject)
            mail_body = urllib.parse.quote(report_text)

            if st.button("Prepare Mail", use_container_width=True):
                if not to_email:
                    st.warning("Enter recipient email first")
                else:
                    to_encoded = urllib.parse.quote(to_email)
                    gmail_url = f"https://mail.google.com/mail/?view=cm&fs=1&to={to_encoded}&su={mail_subject}&body={mail_body}"
                    mailto_url = f"mailto:{to_email}?subject={mail_subject}&body={mail_body}"

                    st.success("Mail template ready, click below to send")
                    st.link_button("Open in Gmail (web)", gmail_url, use_container_width=True)
                    st.link_button("Open in Mail App", mailto_url, use_container_width=True)
                    st.caption("This just opens the mail already filled in, using your own logged-in Gmail / mail app. We don't touch your password.")

    with colB:
        with st.container(border=True):
            st.markdown('<div class="section-title">🟢 Send via WhatsApp</div>', unsafe_allow_html=True)
            wa_text = urllib.parse.quote(report_text)
            wa_url = "https://wa.me/?text=" + wa_text
            st.link_button("Share on WhatsApp", wa_url, use_container_width=True)
            st.caption("Opens WhatsApp, pick a contact and send")

# TODO: maybe add monthly comparison graph later
# TODO: pdf export instead of just csv
